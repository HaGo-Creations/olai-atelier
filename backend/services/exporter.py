"""Export resources to DOCX and PDF with Brand & Layout applied.

DOCX:  python-docx (headers, footers, logo, typography)
PDF:   reportlab (programmatic layout with branding)

Both formats apply:
  - Logo + school name + address in header (alignable)
  - Footer text or phone/email
  - Heading + Body typography from Brand & Layout
  - Per-document header/footer override from Studio Output
  - Tamil + other Indic scripts via Noto fonts (installed via Dockerfile)
"""

import io
import os
import re
import uuid
from typing import Optional

from models import ExportRequest, Branding, PageLayout, HeaderFooterOverride

EXPORT_DIR = os.getenv("EXPORT_DIR", "/tmp/gemma_exports")
os.makedirs(EXPORT_DIR, exist_ok=True)


# ── Unicode font support ───────────────────────────────────────────────────
#
# reportlab ships with PDF base-14 fonts (Helvetica, Times, Courier) that have
# NO glyphs for Tamil, Hindi, Bengali or any Indic script — any non-Latin
# character renders as a black box or empty rectangle.
#
# To fix this we register Noto Sans (Latin + a wide Unicode coverage) and
# Noto Sans Tamil (the Tamil block specifically) at startup, then use them
# as the default body/heading faces. python-docx doesn't have this problem
# because Word resolves fonts at open time on the user's machine, but we
# also set run.font.complex_script so Word picks an appropriate Indic
# fallback (e.g. Latha / Nirmala UI) for Tamil runs.

# Candidate font paths on Debian / Ubuntu (HF Spaces uses python:slim which
# is Debian-based). We probe each path and use the first match.
_FONT_CANDIDATES = {
    "body": [
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ],
    "bold": [
        "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
    ],
    "italic": [
        "/usr/share/fonts/truetype/noto/NotoSans-Italic.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
    ],
    "tamil": [
        "/usr/share/fonts/truetype/noto/NotoSansTamil-Regular.ttf",
        "/usr/share/fonts/truetype/noto/NotoSerifTamil-Regular.ttf",
    ],
    "tamil_bold": [
        "/usr/share/fonts/truetype/noto/NotoSansTamil-Bold.ttf",
        "/usr/share/fonts/truetype/noto/NotoSerifTamil-Bold.ttf",
    ],
    "devanagari": [
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Regular.ttf",
    ],
    "devanagari_bold": [
        "/usr/share/fonts/truetype/noto/NotoSansDevanagari-Bold.ttf",
    ],
}

# Internal font-name registry. Populated lazily on first PDF export.
_FONT_NAMES = {
    "body": "Helvetica",          # safe defaults if Noto isn't installed
    "bold": "Helvetica-Bold",
    "italic": "Helvetica-Oblique",
}
_FONTS_REGISTERED = False


def _first_existing(paths):
    for p in paths:
        if os.path.exists(p):
            return p
    return None


def _register_unicode_fonts():
    """Register Noto fonts with reportlab. Idempotent."""
    global _FONTS_REGISTERED, _FONT_NAMES
    if _FONTS_REGISTERED:
        return

    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
    except ImportError:
        return

    body = _first_existing(_FONT_CANDIDATES["body"])
    bold = _first_existing(_FONT_CANDIDATES["bold"]) or body
    italic = _first_existing(_FONT_CANDIDATES["italic"]) or body

    if body:
        try:
            pdfmetrics.registerFont(TTFont("NotoSans", body))
            pdfmetrics.registerFont(TTFont("NotoSans-Bold", bold))
            pdfmetrics.registerFont(TTFont("NotoSans-Italic", italic))
            # Register the family so <b>/<i> inline tags in Paragraph work.
            registerFontFamily(
                "NotoSans",
                normal="NotoSans", bold="NotoSans-Bold", italic="NotoSans-Italic",
                boldItalic="NotoSans-Bold",
            )
            _FONT_NAMES["body"] = "NotoSans"
            _FONT_NAMES["bold"] = "NotoSans-Bold"
            _FONT_NAMES["italic"] = "NotoSans-Italic"
        except Exception as e:
            print(f"[exporter] Could not register NotoSans: {e}")

    # Register script-specific fonts. reportlab can't do automatic font
    # fallback per glyph, but we can use these explicitly when we detect
    # the script (see _split_by_script below) or pre-substitute them on
    # runs known to be Tamil/Hindi.
    for key, family in [
        ("tamil", "NotoSansTamil"),
        ("devanagari", "NotoSansDevanagari"),
    ]:
        reg = _first_existing(_FONT_CANDIDATES[key])
        bold_path = _first_existing(_FONT_CANDIDATES[f"{key}_bold"]) or reg
        if reg:
            try:
                pdfmetrics.registerFont(TTFont(family, reg))
                pdfmetrics.registerFont(TTFont(f"{family}-Bold", bold_path))
                registerFontFamily(
                    family,
                    normal=family, bold=f"{family}-Bold",
                    italic=family, boldItalic=f"{family}-Bold",
                )
                _FONT_NAMES[key] = family
            except Exception as e:
                print(f"[exporter] Could not register {family}: {e}")

    _FONTS_REGISTERED = True
    print(f"[exporter] Fonts registered: {_FONT_NAMES}")


# Unicode ranges per script. Used to wrap script-specific runs in their own
# <font> tag inside Paragraph markup so reportlab uses the right TTF.
_SCRIPT_RANGES = {
    "tamil":      (0x0B80, 0x0BFF),
    "devanagari": (0x0900, 0x097F),
    # Other ranges can be added as more Noto fonts are bundled.
}


def _wrap_unicode_runs(text: str) -> str:
    """Wrap Tamil / Devanagari sub-strings in reportlab <font> tags so each
    run renders with its script-appropriate TTF.

    Latin text falls through to the paragraph's default font.
    """
    if not text:
        return text

    out = []
    current_script = None  # "tamil" | "devanagari" | None
    buf = []

    def flush():
        if not buf:
            return
        chunk = "".join(buf)
        if current_script and current_script in _FONT_NAMES:
            out.append(f'<font name="{_FONT_NAMES[current_script]}">{chunk}</font>')
        else:
            out.append(chunk)
        buf.clear()

    for ch in text:
        cp = ord(ch)
        script = None
        for name, (lo, hi) in _SCRIPT_RANGES.items():
            if lo <= cp <= hi:
                script = name
                break
        if script != current_script:
            flush()
            current_script = script
        buf.append(ch)
    flush()
    return "".join(out)



async def export_resource(req: ExportRequest) -> str:
    """Export to DOCX or PDF, return file path."""
    safe_title = re.sub(r"[^\w\s-]", "", req.title).strip().replace(" ", "_") or "resource"
    short_id = uuid.uuid4().hex[:8]

    if req.format == "docx":
        path = os.path.join(EXPORT_DIR, f"{safe_title}_{short_id}.docx")
        _export_docx(req, path)
    elif req.format == "pdf":
        path = os.path.join(EXPORT_DIR, f"{safe_title}_{short_id}.pdf")
        _export_pdf(req, path)
    else:
        raise ValueError(f"Unsupported format: {req.format}. Use docx or pdf.")

    return path


# ── DOCX ───────────────────────────────────────────────────────────────────

def _export_docx(req: ExportRequest, path: str) -> None:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    branding = req.branding
    layout = (branding.layouts.get("docx") if branding else None) or PageLayout()
    override = req.header_footer_override

    # Page setup
    section = doc.sections[0]
    section.top_margin = Mm(layout.margin_top)
    section.bottom_margin = Mm(layout.margin_bottom)
    section.left_margin = Mm(layout.margin_left)
    section.right_margin = Mm(layout.margin_right)

    # Header
    if branding and branding.apply_on_export and override.include_header:
        _add_docx_header(section, branding, layout)

    # Footer
    if branding and branding.apply_on_export and override.include_footer:
        _add_docx_footer(section, branding, layout)

    # Body — parse Markdown into paragraphs
    _add_docx_body(doc, req.content_markdown, layout)

    doc.save(path)


def _set_run_unicode_font(run, text: str):
    """Set font names on a docx run so Word renders Indic glyphs correctly.

    Word uses three font slots per run: `ascii` for Latin, `eastAsia` for
    CJK, and `cs` (complex script) for scripts like Tamil / Arabic /
    Devanagari. python-docx only exposes the ascii slot via `run.font.name`,
    so we drop down to the underlying XML to set the others.

    We always set:
      - ascii  = Calibri (works on every Word install)
      - cs     = Nirmala UI (Windows default Indic) with Latha (older Tamil-
                 only fallback) listed via altName so older systems also work
    """
    from docx.oxml.ns import qn
    run.font.name = "Calibri"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Calibri")
    rFonts.set(qn("w:hAnsi"), "Calibri")
    # Nirmala UI is on Windows by default and covers all Indic scripts.
    # macOS/LibreOffice will substitute (e.g. macOS picks the system
    # Indic-capable font automatically when cs is set).
    rFonts.set(qn("w:cs"), "Nirmala UI")
    rFonts.set(qn("w:eastAsia"), "Calibri")
    # Flag complex script so Word uses the cs font for the relevant runs.
    if _contains_complex_script(text):
        cs_el = OxmlElement("w:cs") if False else rPr.find(qn("w:cs"))
        # The <w:cs/> toggle element marks the run as complex-script.
        from docx.oxml import OxmlElement as OE
        if rPr.find(qn("w:cs")) is None:
            rPr.append(OE("w:cs"))


def _contains_complex_script(text: str) -> bool:
    """True if any char is in a complex-script Unicode range we care about."""
    for ch in text or "":
        cp = ord(ch)
        for lo, hi in _SCRIPT_RANGES.values():
            if lo <= cp <= hi:
                return True
    return False


def _add_run(paragraph, text: str, *, bold: bool = False, size_pt=None):
    """Add a run with Unicode-safe font configuration."""
    from docx.shared import Pt
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    _set_run_unicode_font(run, text)
    return run


def _add_docx_header(section, branding: Branding, layout: PageLayout):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    header = section.header
    p = header.paragraphs[0]
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                 "center": WD_ALIGN_PARAGRAPH.CENTER,
                 "right": WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = align_map.get(layout.logo_position, WD_ALIGN_PARAGRAPH.LEFT)

    if branding.school_name:
        _add_run(p, branding.school_name + "\n", bold=True,
                 size_pt=layout.heading_size * 0.55)
    if branding.address:
        _add_run(p, branding.address + "\n",
                 size_pt=layout.body_size * 0.8)


def _add_docx_footer(section, branding: Branding, layout: PageLayout):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    footer = section.footer
    p = footer.paragraphs[0]
    align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT,
                 "center": WD_ALIGN_PARAGRAPH.CENTER,
                 "right": WD_ALIGN_PARAGRAPH.RIGHT}
    p.alignment = align_map.get(layout.footer_alignment, WD_ALIGN_PARAGRAPH.CENTER)

    text = branding.footer_text or " ".join(filter(None, [branding.phone, branding.email]))
    _add_run(p, text, size_pt=max(8, layout.body_size * 0.75))


def _add_docx_body(doc, markdown: str, layout):
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.size = Pt(layout.body_size)
    # Set the Normal style's CS font so any non-_add_run paragraphs we create
    # (e.g. via doc.add_paragraph for bullets) inherit Indic-capable shaping.
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts")) or (rPr.append(OxmlElement("w:rFonts")) or rPr.find(qn("w:rFonts")))
        rFonts.set(qn("w:cs"), "Nirmala UI")
    except Exception:
        pass

    for raw_line in markdown.split("\n"):
        line = raw_line.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            _add_run(p, line[2:], bold=True, size_pt=layout.heading_size)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            _add_run(p, line[3:], bold=True, size_pt=layout.heading_size * 0.85)
        elif line.startswith("### "):
            p = doc.add_paragraph()
            _add_run(p, line[4:], bold=True, size_pt=layout.heading_size * 0.75)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_run(p, line[2:])
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_run(p, re.sub(r"^\d+\.\s", "", line))
        elif line == "---":
            p = doc.add_paragraph()
            _add_run(p, "─" * 40)
        else:
            p = doc.add_paragraph()
            _add_run(p, line)


# ── PDF ────────────────────────────────────────────────────────────────────

def _export_pdf(req: ExportRequest, path: str):
    from reportlab.lib.pagesizes import A4, LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, PageBreak, HRFlowable,
    )

    # Make sure Tamil + Devanagari + Latin fonts are available before we
    # build any flowables. Idempotent — safe to call on every export.
    _register_unicode_fonts()

    branding = req.branding
    layout = (branding.layouts.get("pdf") if branding else None) or PageLayout()
    override = req.header_footer_override

    page_size = A4 if layout.page_size == "a4" else LETTER

    doc = SimpleDocTemplate(
        path, pagesize=page_size,
        topMargin=layout.margin_top * mm,
        bottomMargin=layout.margin_bottom * mm,
        leftMargin=layout.margin_left * mm,
        rightMargin=layout.margin_right * mm,
        title=req.title,
    )

    styles = getSampleStyleSheet()
    h_style = ParagraphStyle(
        "H", parent=styles["Heading1"],
        fontName=_FONT_NAMES["bold"],
        fontSize=layout.heading_size,
        leading=layout.heading_size * 1.3,
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontName=_FONT_NAMES["bold"],
        fontSize=layout.heading_size * 0.85,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["BodyText"],
        fontName=_FONT_NAMES["body"],
        fontSize=layout.body_size,
        leading=layout.body_size * layout.line_spacing,
    )

    flowables = []

    for raw_line in req.content_markdown.split("\n"):
        line = raw_line.rstrip()
        if not line:
            flowables.append(Spacer(1, layout.body_size))
            continue
        if line.startswith("# "):
            flowables.append(Paragraph(_wrap_unicode_runs(_html_escape(line[2:])), h_style))
        elif line.startswith("## "):
            flowables.append(Paragraph(_wrap_unicode_runs(_html_escape(line[3:])), h2_style))
        elif line == "---":
            flowables.append(HRFlowable(width="100%"))
        else:
            flowables.append(Paragraph(_wrap_unicode_runs(_html_escape(line)), body_style))

    def _on_page(canvas, doc_):
        if branding and branding.apply_on_export and override.include_header:
            _draw_pdf_header(canvas, doc_, branding, layout)
        if branding and branding.apply_on_export and override.include_footer:
            _draw_pdf_footer(canvas, doc_, branding, layout)

    doc.build(flowables, onFirstPage=_on_page, onLaterPages=_on_page)


def _pick_font_for(text: str, bold: bool = False) -> str:
    """Return the font name best suited to render the dominant script in
    text. Heuristic: first detected script wins. Falls back to body."""
    for ch in text:
        cp = ord(ch)
        for name, (lo, hi) in _SCRIPT_RANGES.items():
            if lo <= cp <= hi and name in _FONT_NAMES:
                # Tamil/Devanagari registered families always have a -Bold
                # variant when registered. Use it for bold contexts.
                return f"{_FONT_NAMES[name]}-Bold" if bold else _FONT_NAMES[name]
    return _FONT_NAMES["bold"] if bold else _FONT_NAMES["body"]


def _draw_pdf_header(canvas, doc_, branding: Branding, layout: PageLayout):
    from reportlab.lib.units import mm
    canvas.saveState()
    y = doc_.pagesize[1] - 12 * mm
    if layout.logo_position == "left":
        x = 20 * mm
    elif layout.logo_position == "right":
        x = doc_.pagesize[0] - 20 * mm
    else:
        x = doc_.pagesize[0] / 2

    if branding.school_name:
        canvas.setFont(_pick_font_for(branding.school_name, bold=True),
                       layout.heading_size * 0.55)
        if layout.logo_position == "right":
            canvas.drawRightString(x, y, branding.school_name)
        elif layout.logo_position == "center":
            canvas.drawCentredString(x, y, branding.school_name)
        else:
            canvas.drawString(x, y, branding.school_name)
    if branding.address:
        canvas.setFont(_pick_font_for(branding.address), layout.body_size * 0.7)
        if layout.logo_position == "right":
            canvas.drawRightString(x, y - 10, branding.address)
        elif layout.logo_position == "center":
            canvas.drawCentredString(x, y - 10, branding.address)
        else:
            canvas.drawString(x, y - 10, branding.address)
    canvas.restoreState()


def _draw_pdf_footer(canvas, doc_, branding: Branding, layout: PageLayout):
    from reportlab.lib.units import mm
    canvas.saveState()
    text = branding.footer_text or " ".join(filter(None, [branding.phone, branding.email]))
    if not text:
        canvas.restoreState()
        return
    canvas.setFont(_pick_font_for(text), layout.body_size * 0.7)
    y = 10 * mm
    if layout.footer_alignment == "left":
        canvas.drawString(20 * mm, y, text)
    elif layout.footer_alignment == "right":
        canvas.drawRightString(doc_.pagesize[0] - 20 * mm, y, text)
    else:
        canvas.drawCentredString(doc_.pagesize[0] / 2, y, text)
    canvas.restoreState()


def _html_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")